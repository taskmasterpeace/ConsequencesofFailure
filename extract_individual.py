import os
from docx import Document

def extract_text_from_docx(file_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(file_path)
        full_text = []

        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e:
        return f"ERROR reading file: {str(e)}"

# Files to extract in priority order
files = [
    "Liu_Xiao_Story_Spine.docx",
    "Establishment_24.docx",
    "COUNTRIES_SEASONS_AI.docx",
    "Atlas_Protocol_Episodes.docx",
    "The_Atlas_Protocol_notes.docx",
    "Kaiser_Main.docx",
    "Eziobi_Story_1.docx",
    "AI_Cult_Story.docx",
    "SPEAR_Pilot_Script.docx",
    "Character_Creation.docx"
]

base_path = r"C:\git\COF\COF_Stories"
output_dir = r"C:\git\COF\extracted_files"

# Create output directory if it doesn't exist
os.makedirs(output_dir, exist_ok=True)

for filename in files:
    file_path = os.path.join(base_path, filename)
    if os.path.exists(file_path):
        text = extract_text_from_docx(file_path)

        # Save to individual text file
        output_filename = filename.replace('.docx', '.txt')
        output_path = os.path.join(output_dir, output_filename)

        with open(output_path, 'w', encoding='utf-8') as out_file:
            out_file.write(text)

        print(f"Extracted: {filename} -> {output_filename}")
    else:
        print(f"NOT FOUND: {filename}")

print(f"\nExtraction complete! Files saved to: {output_dir}")
